from docx import Document
from docx.shared import Pt

d = Document()
d.add_paragraph("Hello ").add_run("Python").bold = True
d.save("1.docx")


doc2 = Document()
new_paragraph = doc2.add_paragraph()

doc = Document("1.docx")
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            print(run.text)
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = True
            new_run.font.size = Pt(16)
            new_run.font.name = 'Times New Roman'
doc2.save("2.docx")


